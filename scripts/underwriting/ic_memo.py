"""
ic_memo.py — Institutional IC memo docx generator.

Produces a standard acquisition IC memo from any engine's underwriting output
(multifamily / commercial / hospitality). Uses the shared docx_style helpers
(vendor/asset-management) for consistent institutional appearance.

CLI:
    python -m scripts.underwriting.ic_memo deal.yaml [-o memo.docx]

Auto-detects asset class via `property.asset_class` in the YAML and dispatches
to the correct engine. Section order:

    Cover page
    1. Executive Summary
    2. Property Overview
    3. Transaction Structure
    4. Operating Pro Forma
    5. Debt
    6. Returns & Waterfall
    7. Exit
    8. Risks (analyst-fill stubs)
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

import scripts  # noqa: F401
from docx import Document
from docx_style import (
    add_bullets,
    add_cover_page,
    add_heading,
    add_para,
    add_source,
    add_table,
    apply_memo_styles,
)

from .excel_summary import SummaryPayload, build_payload


# ---------------------------------------------------------------------------
# Memo-specific payload extension
# ---------------------------------------------------------------------------

@dataclass
class MemoYearLine:
    year: int
    revenue: float
    opex: float
    noi: float
    ncf_unlevered: float
    ncf_levered: float


@dataclass
class CommercialMemoBlock:
    """Commercial-specific tables: rent roll, WALT, concentration, rollover."""
    walt_yrs: float                              # weighted-average lease term (by rent)
    rent_roll: list[tuple[str, str, int, float, float, str, str, float]]
    # (tenant, suite, sf, base_rent_psf, annual_rent, lease_type, lease_end_iso, pct_of_rent)
    top_tenants: list[tuple[str, int, float, float]]
    # (tenant, sf, annual_rent, pct_of_rent) -- top 5 by rent
    rollover: list[tuple[int, int, float, float, float, float]]
    # (year, sf_rolling, pct_of_rba, in_place_rent, market_rent_at_roll, mtm_pct)


@dataclass
class DCMemoBlock:
    """Data-center-specific tables + negotiation playbook.

    `kind` selects which leasing-tactics subset is rendered. `tenancy_rows` is a
    free-form list of (label, value) pairs for the tenancy snapshot table --
    wholesale shows MW utilization + per-contract; colo shows cabinet mix +
    occupancy ramp.
    """
    kind: Literal["wholesale", "colo"]
    tenancy_rows: list[tuple[str, str]]
    rollover: list[tuple[int, float, float, float, float]] | None
    # wholesale only: (year, mw_rolling, in_place_rent, market_rent_at_roll, mtm_pct)
    top_contracts: list[tuple[str, float, float, float, str, str]] | None
    # wholesale only: (tenant, mw, $/kW/mo, annual_rent, lease_end_iso, pass_through)
    cabinet_mix: list[tuple[str, int, float, float, float, float]] | None
    # colo only: (name, count, kw/cab, total_kw, in_place_mrr, market_mrr)


@dataclass
class InfraMemoBlock:
    """Infrastructure-specific tables: generation profile, revenue stream roster,
    contracted-vs-merchant mix by year, and tax credits.
    """
    summary_rows: list[tuple[str, str]]
    # (label, value): technology, market, nameplate, CF, realized CF, net gen Yr1, # streams, Yr-1 contracted share
    stream_rows: list[tuple[str, str, str, str, str, str, float]]
    # (label, kind, counterparty (+rating), term, headline_rate, allotment, yr1_revenue)
    mix_rows: list[tuple[int, float, float, float, float, float, float]]
    # (year, ppa_rev, avail_rev, merch_rev, ptc_rev, total_rev, contracted_share)
    tax_credit_rows: list[tuple[str, str]]
    # (label, value): ITC $, PTC $/MWh, PTC term


# ---------------------------------------------------------------------------
# Formatters
# ---------------------------------------------------------------------------

def _fmt_dollar(v: float) -> str:
    return f"${v:,.0f}"


def _fmt_dollar_m(v: float) -> str:
    return f"${v / 1_000_000:,.2f}M"


def _fmt_pct(v: float, digits: int = 2) -> str:
    return f"{v * 100:.{digits}f}%"


def _fmt_multiple(v: float) -> str:
    return f"{v:.2f}x"


def _fmt_per_denom(v: float, label: str) -> str:
    return f"${v:,.0f} {label}"


# ---------------------------------------------------------------------------
# Writer
# ---------------------------------------------------------------------------

def write_ic_memo(
    p: SummaryPayload,
    year_lines: list[MemoYearLine],
    out_path: Path,
    revenue_label: str = "Revenue",
    preparer: str = "Acquisitions",
    commercial: CommercialMemoBlock | None = None,
    datacenter: DCMemoBlock | None = None,
    infrastructure: InfraMemoBlock | None = None,
) -> Path:
    """Render the IC memo as a docx at `out_path`."""
    doc = Document()
    apply_memo_styles(doc, theme="modern")

    add_cover_page(
        doc,
        title=f"Acquisition Memorandum  —  {p.deal_name}",
        recipient="Investment Committee",
        preparer=preparer,
        date_str=p.close_date.strftime("%B %d, %Y"),
        confidentiality=True,
    )

    # ---- 1. Executive Summary ----
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(
        doc,
        f"{p.sponsor} proposes the acquisition of {p.property_name}, a "
        f"{p.asset_class.lower()} asset located in {p.submarket} "
        f"({p.address}), for a purchase price of {_fmt_dollar(p.purchase_price)} "
        f"({_fmt_per_denom(p.purchase_price / p.denom_value, p.per_denom_label)}). "
        f"Total all-in basis of {_fmt_dollar(p.all_in_basis)} "
        f"({_fmt_per_denom(p.all_in_basis / p.denom_value, p.per_denom_label)}) "
        f"underwrites to {_fmt_pct(p.going_in_cap)} going-in cap, "
        f"{_fmt_pct(p.roc_trended)} stabilized YoC on all-in basis "
        f"(Yr {p.stab_yr}), on a {p.hold_yrs}-year hold."
    )
    add_para(
        doc,
        f"Projected returns: project IRR {_fmt_pct(p.project_irr)} / MOIC "
        f"{_fmt_multiple(p.project_moic)}; LP net IRR "
        f"{_fmt_pct(p.lp_irr)} / MOIC {_fmt_multiple(p.lp_moic)}; "
        f"GP IRR {_fmt_pct(p.gp_irr)} / MOIC {_fmt_multiple(p.gp_moic)} "
        f"(inclusive of {_fmt_pct(p.gp_coinvest_pct, 0)} co-invest)."
    )

    add_table(
        doc,
        headers=["Metric", "Value"],
        rows=[
            ["Purchase Price",       _fmt_dollar(p.purchase_price)],
            [f"Price {p.per_denom_label}", _fmt_dollar(p.purchase_price / p.denom_value)],
            ["All-In Basis",         _fmt_dollar(p.all_in_basis)],
            ["Going-In Cap",                          _fmt_pct(p.going_in_cap)],
            [f"Stab Cap on Price (Yr {p.stab_yr})",   _fmt_pct(p.stabilized_cap)],
            [f"YoC on All-In (Yr {p.stab_yr})",       _fmt_pct(p.roc_trended)],
            ["Exit Cap",                              _fmt_pct(p.exit_cap)],
            ["Untrended YOC @ Stab", _fmt_pct(p.roc_untrended)],
            ["Trended YOC @ Stab",   _fmt_pct(p.roc_trended)],
            ["YOC @ Exit (FTM)",     _fmt_pct(p.roc_exit_ftm)],
            ["LTV",                  _fmt_pct(p.ltv, 1)],
            ["Yr 1 DSCR",            _fmt_multiple(p.dscr)],
            ["Project IRR",          _fmt_pct(p.project_irr)],
            ["LP Net IRR",           _fmt_pct(p.lp_irr)],
            ["GP IRR",               _fmt_pct(p.gp_irr)],
        ],
        numeric_cols=[1],
    )

    # ---- 2. Property Overview ----
    add_heading(doc, "2. Property Overview", level=1)
    add_table(
        doc,
        headers=["Attribute", "Detail"],
        rows=[
            ["Property",        p.property_name],
            ["Address",         p.address],
            ["Submarket",       p.submarket],
            ["Asset Class",     p.asset_class],
            [p.denom_label,     f"{p.denom_value:,.0f}"],
            ["Close Date",      p.close_date.isoformat()],
            ["Hold (yrs)",      str(p.hold_yrs)],
        ],
        numeric_cols=[],
    )

    # ---- 2b. Tenancy & Rollover (commercial only) ----
    if commercial is not None:
        add_heading(doc, "2a. Tenancy & Rollover", level=1)
        add_para(
            doc,
            f"Weighted-average lease term (WALT, rent-weighted): "
            f"{commercial.walt_yrs:.1f} years."
        )
        add_para(doc, "Top tenants by annual rent:")
        add_table(
            doc,
            headers=["Tenant", "SF", "Annual Rent", "% of Rent"],
            rows=[
                [t, f"{sf:,}", _fmt_dollar(rent), _fmt_pct(pct, 1)]
                for (t, sf, rent, pct) in commercial.top_tenants
            ],
            numeric_cols=[1, 2, 3],
        )
        add_para(doc, "Rent roll detail:")
        add_table(
            doc,
            headers=["Tenant", "Suite", "SF", "$/SF", "Annual Rent", "Type", "Expires", "% of Rent"],
            rows=[
                [t, s, f"{sf:,}", f"${psf:.2f}", _fmt_dollar(annual),
                 typ, exp, _fmt_pct(pct, 1)]
                for (t, s, sf, psf, annual, typ, exp, pct) in commercial.rent_roll
            ],
            numeric_cols=[2, 3, 4, 7],
        )
        add_para(doc, "Rollover schedule (lease expirations during hold):")
        add_table(
            doc,
            headers=["Year", "SF Rolling", "% of RBA", "In-Place Rent",
                     "Market Rent at Roll", "MTM Spread"],
            rows=[
                [f"Yr {yr}", f"{sf:,}", _fmt_pct(rba_pct, 1),
                 _fmt_dollar(ip), _fmt_dollar(mk), _fmt_pct(mtm, 1)]
                for (yr, sf, rba_pct, ip, mk, mtm) in commercial.rollover
            ],
            numeric_cols=[1, 2, 3, 4, 5],
        )

    # ---- 2b. Data Center tenancy / capacity ----
    if datacenter is not None:
        heading = (
            "2a. Tenancy & Power Capacity" if datacenter.kind == "wholesale"
            else "2a. Cabinet Mix & Lease-Up"
        )
        add_heading(doc, heading, level=1)
        add_table(
            doc,
            headers=["Item", "Value"],
            rows=[list(row) for row in datacenter.tenancy_rows],
            numeric_cols=[1],
        )
        if datacenter.kind == "wholesale" and datacenter.top_contracts:
            add_para(doc, "Contract roster (top contracts by MW):")
            add_table(
                doc,
                headers=["Tenant", "MW", "$/kW/mo", "Annual Rent", "Lease End", "Pass-Thru"],
                rows=[
                    [t, f"{mw:.2f}", f"${rate:.2f}", _fmt_dollar(ann), end, pt]
                    for (t, mw, rate, ann, end, pt) in datacenter.top_contracts
                ],
                numeric_cols=[1, 2, 3],
            )
            if datacenter.rollover:
                add_para(doc, "Rollover schedule (lease expirations during hold):")
                add_table(
                    doc,
                    headers=["Year", "MW Rolling", "In-Place Rent",
                             "Market Rent at Roll", "MTM Spread"],
                    rows=[
                        [f"Yr {yr}", f"{mw:.2f}", _fmt_dollar(ip),
                         _fmt_dollar(mk), _fmt_pct(mtm, 1)]
                        for (yr, mw, ip, mk, mtm) in datacenter.rollover
                    ],
                    numeric_cols=[1, 2, 3, 4],
                )
        if datacenter.kind == "colo" and datacenter.cabinet_mix:
            add_para(doc, "Cabinet inventory (in-place vs. market MRR):")
            add_table(
                doc,
                headers=["Cabinet Type", "Count", "kW / Cab", "Total kW",
                         "In-Place MRR", "Market MRR"],
                rows=[
                    [n, f"{c:,}", f"{kw:.1f}", f"{tkw:,.0f}",
                     _fmt_dollar(ip), _fmt_dollar(mk)]
                    for (n, c, kw, tkw, ip, mk) in datacenter.cabinet_mix
                ],
                numeric_cols=[1, 2, 3, 4, 5],
            )

    # ---- 2a. Infrastructure: Generation & Revenue Mix ----
    if infrastructure is not None:
        add_heading(doc, "2a. Generation & Revenue Mix", level=1)
        add_table(
            doc,
            headers=["Item", "Value"],
            rows=[list(row) for row in infrastructure.summary_rows],
            numeric_cols=[1],
        )
        add_para(doc, "Revenue stream roster:")
        add_table(
            doc,
            headers=["Stream", "Type", "Counterparty", "Term",
                     "Headline Rate", "Allotment", "Yr 1 Revenue"],
            rows=[
                [lbl, kind, cpty, term, rate, allot, _fmt_dollar(rev)]
                for (lbl, kind, cpty, term, rate, allot, rev)
                in infrastructure.stream_rows
            ],
            numeric_cols=[6],
        )
        add_para(doc, "Contracted vs. merchant revenue by hold year:")
        add_table(
            doc,
            headers=["Year", "PPA", "Availability", "Merchant", "PTC",
                     "Total", "Contracted %"],
            rows=[
                [f"Yr {yr}", _fmt_dollar(ppa), _fmt_dollar(av),
                 _fmt_dollar(mc), _fmt_dollar(ptc), _fmt_dollar(tot),
                 _fmt_pct(cs, 1)]
                for (yr, ppa, av, mc, ptc, tot, cs) in infrastructure.mix_rows
            ],
            numeric_cols=[1, 2, 3, 4, 5, 6],
        )
        if infrastructure.tax_credit_rows:
            add_para(doc, "Federal tax credits (monetized as cash):")
            add_table(
                doc,
                headers=["Item", "Value"],
                rows=[list(row) for row in infrastructure.tax_credit_rows],
                numeric_cols=[1],
            )

    # ---- 3. Transaction Structure ----
    add_heading(doc, "3. Transaction Structure", level=1)
    add_para(doc, "Sources & uses:")
    uses = [
        ["Purchase Price",        _fmt_dollar(p.purchase_price)],
        ["Closing Costs",         _fmt_dollar(p.closing_costs)],
        ["Initial CapEx",         _fmt_dollar(p.initial_capex)],
        ["Value-Add / PIP",       _fmt_dollar(p.value_add_capex_total)],
        ["Day-One Reserves",      _fmt_dollar(p.day_one_reserves)],
        ["All-In Basis",          _fmt_dollar(p.all_in_basis)],
    ]
    add_table(doc, headers=["Use", "Amount"], rows=uses, numeric_cols=[1], total_row=True)
    add_para(doc, "Capitalization:")
    add_table(
        doc,
        headers=["Source", "Amount", "% of Capitalization"],
        rows=[
            ["Senior Debt",     _fmt_dollar(p.loan_amount),
             _fmt_pct(p.loan_amount / p.all_in_basis, 1)],
            ["LP Equity",       _fmt_dollar(p.lp_contributed),
             _fmt_pct(p.lp_contributed / p.all_in_basis, 1)],
            ["GP Co-Invest",    _fmt_dollar(p.gp_contributed),
             _fmt_pct(p.gp_contributed / p.all_in_basis, 1)],
        ],
        numeric_cols=[1, 2],
    )

    # ---- 4. Operating Pro Forma ----
    add_heading(doc, "4. Operating Pro Forma", level=1)
    rows = [
        [
            f"Yr {yl.year}",
            _fmt_dollar(yl.revenue),
            _fmt_dollar(-yl.opex),
            _fmt_dollar(yl.noi),
            _fmt_dollar(yl.ncf_unlevered),
            _fmt_dollar(yl.ncf_levered),
        ]
        for yl in year_lines
    ]
    add_table(
        doc,
        headers=["Period", revenue_label, "OpEx", "NOI", "Unlevered NCF", "Levered NCF"],
        rows=rows,
        numeric_cols=[1, 2, 3, 4, 5],
    )

    # ---- 5. Debt ----
    add_heading(doc, "5. Debt", level=1)
    add_table(
        doc,
        headers=["Term", "Value"],
        rows=[
            ["Loan Amount",        _fmt_dollar(p.loan_amount)],
            ["Binding Constraint", p.binding_constraint],
            ["LTV",                _fmt_pct(p.ltv, 1)],
            ["Yr 1 DSCR",          _fmt_multiple(p.dscr)],
            ["Yr 1 Debt Yield",    _fmt_pct(p.debt_yield)],
            ["All-In Rate",        _fmt_pct(p.rate)],
            ["Term",               f"{p.term_yrs} years"],
            ["Amortization",       f"{p.amort_yrs} years"],
            ["IO Period",          f"{p.io_period_yrs} years"],
        ],
        numeric_cols=[1],
    )

    # ---- 6. Returns & Waterfall ----
    add_heading(doc, "6. Returns & Waterfall", level=1)
    add_para(doc, "Projected returns:")
    add_table(
        doc,
        headers=["Party", "IRR", "MOIC", "Contributed", "Distributed"],
        rows=[
            ["Project (Total Equity)",
             _fmt_pct(p.project_irr), _fmt_multiple(p.project_moic),
             _fmt_dollar(p.lp_contributed + p.gp_contributed),
             _fmt_dollar(p.lp_distributed + p.gp_distributed)],
            ["LP (net of waterfall)",
             _fmt_pct(p.lp_irr), _fmt_multiple(p.lp_moic),
             _fmt_dollar(p.lp_contributed), _fmt_dollar(p.lp_distributed)],
            [f"GP (co-invest {_fmt_pct(p.gp_coinvest_pct, 0)} + promote)",
             _fmt_pct(p.gp_irr), _fmt_multiple(p.gp_moic),
             _fmt_dollar(p.gp_contributed), _fmt_dollar(p.gp_distributed)],
        ],
        numeric_cols=[1, 2, 3, 4],
    )

    add_para(doc, "Promote structure (multi-tier IRR hurdle):")
    tier_rows = []
    for i, t in enumerate(p.tiers):
        is_residual = (i == len(p.tiers) - 1)
        tier_rows.append([
            t.label or f"Tier {i+1}",
            "Residual" if is_residual else _fmt_pct(t.hurdle_irr, 1),
            _fmt_pct(t.promote_pct, 1),
            _fmt_dollar(t.lp_total),
            _fmt_dollar(t.gp_total),
        ])
    add_table(
        doc,
        headers=["Tier", "Hurdle IRR", "GP Promote", "LP Cash", "GP Promote $"],
        rows=tier_rows,
        numeric_cols=[1, 2, 3, 4],
    )

    # ---- 7. Exit ----
    add_heading(doc, "7. Exit", level=1)
    add_table(
        doc,
        headers=["Item", "Value"],
        rows=[
            ["Exit Year",                  f"Yr {p.exit_year}"],
            ["Exit NOI Basis",             p.exit_noi_basis],
            ["Exit NOI",                   _fmt_dollar(p.exit_noi)],
            ["Exit Cap",                   _fmt_pct(p.exit_cap)],
            ["Gross Sale Price",           _fmt_dollar(p.gross_sale)],
            ["Cost of Sale",               _fmt_dollar(p.cost_of_sale)],
            ["Loan Payoff",                _fmt_dollar(p.loan_payoff)],
            ["Net Proceeds to Equity",     _fmt_dollar(p.net_proceeds)],
        ],
        numeric_cols=[1],
        total_row=True,
    )

    # ---- 8. Risks ----
    add_heading(doc, "8. Risks & Mitigants", level=1)
    add_para(doc, "Analyst-fill section. Standard risk categories:")
    add_bullets(doc, [
        "Market risk — submarket rent / occupancy / cap rate sensitivity.",
        "Execution risk — value-add or PIP delivery, lease-up pace, downtime.",
        "Capital markets risk — refinance / take-out availability at exit.",
        "Operating risk — expense inflation, tax reassessment, insurance hardening.",
        "Sponsor / counterparty risk — operator quality, GP track record.",
    ])

    # ---- 9. Negotiation Playbook (DC only) ----
    if datacenter is not None:
        from .datacenter.negotiation import (
            acquisition_table_rows,
            leasing_table_rows,
        )
        add_heading(doc, "9. Negotiation Playbook", level=1)
        add_para(
            doc,
            "Institutional negotiation tactics tailored for data center "
            "transactions. The acquisition catalog covers LOI / PSA / closing "
            "levers; the leasing catalog covers post-close lease execution "
            f"({datacenter.kind})."
        )
        add_heading(doc, "9a. Acquisition Tactics", level=2)
        add_table(
            doc,
            headers=["Tactic", "Category", "Lever", "Typical Range", "Rationale"],
            rows=acquisition_table_rows(),
            numeric_cols=[],
        )
        add_heading(doc, "9b. Leasing Tactics", level=2)
        add_table(
            doc,
            headers=["Tactic", "Category", "Lever", "Typical Range", "Rationale"],
            rows=leasing_table_rows(datacenter.kind),
            numeric_cols=[],
        )

    add_source(
        doc,
        "Note: Figures generated by the underwriting engine on "
        f"{p.close_date.isoformat()}; refer to the accompanying workbook for "
        "full assumptions and year-by-year detail."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# Engine adapters
# ---------------------------------------------------------------------------

def _build_mf(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine]]:
    from .models import load_deal
    from .pro_forma import build_pro_forma
    from .waterfall_acq import run_acquisition_waterfall
    deal = load_deal(deal_path)
    pf = build_pro_forma(deal)
    wf = run_acquisition_waterfall(pf)
    value_add_total = (
        deal.capex.value_add_per_unit * deal.property.unit_count
        if deal.capex.value_add_per_unit > 0 else 0.0
    )
    payload = build_payload(
        pf=pf, wf=wf,
        asset_class="Multifamily",
        denom_label="Units",
        denom_value=deal.property.unit_count,
        per_denom_label="/Unit",
        per_denom_fmt="per_unit",
        value_add_capex_total=value_add_total,
    )
    years = [
        MemoYearLine(
            year=yl.year, revenue=yl.egi, opex=yl.total_opex,
            noi=yl.noi, ncf_unlevered=yl.ncf_unlevered, ncf_levered=yl.ncf_levered,
        )
        for yl in pf.years
    ]
    return payload, years


def _build_commercial(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine], CommercialMemoBlock]:
    from .commercial.models import load_commercial_deal
    from .commercial.pro_forma import build_commercial_pro_forma
    from .waterfall_acq import run_acquisition_waterfall
    deal = load_commercial_deal(deal_path)
    pf = build_commercial_pro_forma(deal)
    wf = run_acquisition_waterfall(pf)
    payload = build_payload(
        pf=pf, wf=wf,
        asset_class=deal.property.asset_class.title(),
        denom_label="RBA (SF)",
        denom_value=deal.property.total_rba,
        per_denom_label="/SF",
        per_denom_fmt="per_sf",
        value_add_capex_total=0.0,
    )
    years = [
        MemoYearLine(
            year=yl.year, revenue=yl.egi, opex=yl.total_opex,
            noi=yl.noi, ncf_unlevered=yl.ncf_unlevered, ncf_levered=yl.ncf_levered,
        )
        for yl in pf.years
    ]

    # --- Commercial-specific block ---
    close = deal.acquisition.close_date
    total_in_place = sum(l.base_rent_psf * l.sf for l in deal.property.rent_roll)
    walt_num = sum(
        (l.base_rent_psf * l.sf) * max(0.0, ((l.lease_end - close).days / 365.25))
        for l in deal.property.rent_roll
    )
    walt = walt_num / total_in_place if total_in_place > 0 else 0.0

    rent_roll = []
    for l in deal.property.rent_roll:
        annual = l.base_rent_psf * l.sf
        rent_roll.append((
            l.tenant, l.suite or "-", l.sf, l.base_rent_psf, annual,
            l.lease_type, l.lease_end.isoformat(),
            annual / total_in_place if total_in_place > 0 else 0.0,
        ))

    by_tenant: dict[str, tuple[int, float]] = {}
    for l in deal.property.rent_roll:
        cur_sf, cur_rent = by_tenant.get(l.tenant, (0, 0.0))
        by_tenant[l.tenant] = (cur_sf + l.sf, cur_rent + l.base_rent_psf * l.sf)
    top = sorted(by_tenant.items(), key=lambda kv: -kv[1][1])[:5]
    top_tenants = [
        (t, sf, rent, rent / total_in_place if total_in_place > 0 else 0.0)
        for (t, (sf, rent)) in top
    ]

    rba = deal.property.total_rba
    rollover = [
        (ro.year, ro.sf_rolling, ro.sf_rolling / rba if rba > 0 else 0.0,
         ro.in_place_rent_rolling, ro.market_rent_at_roll, ro.mtm_spread_pct)
        for ro in pf.rollover_schedule
    ]

    block = CommercialMemoBlock(
        walt_yrs=walt, rent_roll=rent_roll, top_tenants=top_tenants, rollover=rollover,
    )
    return payload, years, block


def _build_hospitality(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine]]:
    from .hospitality.models import load_hotel_deal
    from .hospitality.pro_forma import build_hotel_pro_forma
    from .waterfall_acq import run_acquisition_waterfall
    deal = load_hotel_deal(deal_path)
    pf = build_hotel_pro_forma(deal)
    wf = run_acquisition_waterfall(pf)
    payload = build_payload(
        pf=pf, wf=wf,
        asset_class="Hospitality",
        denom_label="Keys",
        denom_value=deal.property.keys,
        per_denom_label="/Key",
        per_denom_fmt="per_unit",
        value_add_capex_total=deal.capex.pip_total,
    )
    years = [
        MemoYearLine(
            year=yl.year,
            revenue=yl.total_revenue,
            opex=yl.total_revenue - yl.noi,   # all costs incl. reserve
            noi=yl.noi,
            ncf_unlevered=yl.ncf_unlevered,
            ncf_levered=yl.ncf_levered,
        )
        for yl in pf.years
    ]
    return payload, years


def _build_dc_wholesale(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine], DCMemoBlock]:
    from .datacenter.models import load_dc_wholesale_deal
    from .datacenter.wholesale_pro_forma import build_wholesale_pro_forma
    from .waterfall_acq import run_acquisition_waterfall
    deal = load_dc_wholesale_deal(deal_path)
    pf = build_wholesale_pro_forma(deal)
    wf = run_acquisition_waterfall(pf)
    payload = build_payload(
        pf=pf, wf=wf,
        asset_class="Data Center (Wholesale)",
        denom_label="Critical MW",
        denom_value=deal.property.mw_critical,
        per_denom_label="/MW",
        per_denom_fmt="dollar",
        value_add_capex_total=0.0,
    )
    years = [
        MemoYearLine(
            year=yl.year, revenue=yl.egi, opex=yl.total_opex,
            noi=yl.noi, ncf_unlevered=yl.ncf_unlevered, ncf_levered=yl.ncf_levered,
        )
        for yl in pf.years
    ]
    prop = deal.property
    tenancy = [
        ("Critical MW",         f"{prop.mw_critical:.2f}"),
        ("Commissioned MW",     f"{prop.mw_commissioned:.2f}"),
        ("Leased MW (in-place)",f"{prop.leased_mw:.2f}"),
        ("Utilization",         _fmt_pct(prop.utilization_pct, 1)),
        ("PUE",                 f"{prop.pue:.2f}"),
        ("Tier Rating",         prop.tier_rating),
        ("In-Place Annual Rent",_fmt_dollar(prop.in_place_annual_rent)),
        ("# Contracts",         f"{len(prop.contracts)}"),
    ]
    top_contracts = sorted(prop.contracts, key=lambda c: -c.mw_leased)[:8]
    top_rows = [
        (c.tenant, c.mw_leased, c.base_rent_kw_mo, c.annual_base_rent,
         c.lease_end.isoformat(), c.power_pass_through)
        for c in top_contracts
    ]
    rollover_rows = [
        (ro.year, ro.mw_rolling, ro.in_place_rent_rolling,
         ro.market_rent_at_roll, ro.mtm_spread_pct)
        for ro in pf.rollover_schedule if ro.mw_rolling > 0
    ]
    block = DCMemoBlock(
        kind="wholesale", tenancy_rows=tenancy,
        rollover=rollover_rows, top_contracts=top_rows, cabinet_mix=None,
    )
    return payload, years, block


def _build_dc_colo(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine], DCMemoBlock]:
    from .datacenter.colo_pro_forma import build_colo_pro_forma
    from .datacenter.models import load_dc_colo_deal
    from .waterfall_acq import run_acquisition_waterfall
    deal = load_dc_colo_deal(deal_path)
    pf = build_colo_pro_forma(deal)
    wf = run_acquisition_waterfall(pf)
    fit_out_total = (
        deal.capex.fit_out_per_cabinet * deal.property.total_cabinets
        if deal.capex.fit_out_per_cabinet > 0 else 0.0
    )
    payload = build_payload(
        pf=pf, wf=wf,
        asset_class="Data Center (Colocation)",
        denom_label="Cabinets",
        denom_value=deal.property.total_cabinets,
        per_denom_label="/Cabinet",
        per_denom_fmt="dollar",
        value_add_capex_total=fit_out_total,
    )
    years = [
        MemoYearLine(
            year=yl.year, revenue=yl.egi,
            opex=yl.total_opex, noi=yl.noi,
            ncf_unlevered=yl.ncf_unlevered, ncf_levered=yl.ncf_levered,
        )
        for yl in pf.years
    ]
    prop = deal.property
    occ_str = " / ".join(f"{o*100:.0f}%" for o in deal.revenue.occupancy[: deal.exit.hold_yrs + 1])
    tenancy = [
        ("Critical MW",        f"{prop.mw_critical:.2f}"),
        ("PUE",                f"{prop.pue:.2f}"),
        ("Tier Rating",        prop.tier_rating),
        ("Total Cabinets",     f"{prop.total_cabinets:,}"),
        ("Total Contracted kW",f"{prop.total_contracted_kw:,.0f}"),
        ("In-Place Gross Rent",_fmt_dollar(prop.in_place_gross_rent)),
        ("Market Gross Rent",  _fmt_dollar(prop.market_gross_rent)),
        ("Occupancy Ramp",     occ_str),
    ]
    cab_rows = [
        (c.name, c.count, c.kw_per_cabinet, c.total_kw, c.in_place_mrr, c.market_mrr)
        for c in prop.cabinet_mix
    ]
    block = DCMemoBlock(
        kind="colo", tenancy_rows=tenancy,
        rollover=None, top_contracts=None, cabinet_mix=cab_rows,
    )
    return payload, years, block


def _build_infrastructure(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine], InfraMemoBlock]:
    from .infrastructure.models import (
        AvailabilityStream,
        MerchantStream,
        PPAStream,
        load_infrastructure_deal,
    )
    from .infrastructure.pro_forma import build_infrastructure_pro_forma
    from .infrastructure.waterfall import run_infrastructure_waterfall
    deal = load_infrastructure_deal(deal_path)
    pf = build_infrastructure_pro_forma(deal)
    wf = run_infrastructure_waterfall(pf)
    gen = deal.property.generation
    payload = build_payload(
        pf=pf, wf=wf,
        asset_class=f"Infrastructure ({gen.technology.title()})",
        denom_label="Nameplate MW (AC)",
        denom_value=gen.nameplate_mw_ac,
        per_denom_label="/MW",
        per_denom_fmt="dollar",
        value_add_capex_total=sum(
            ev.amount for ev in deal.capex.augmentation_schedule
        ),
    )
    years = [
        MemoYearLine(
            year=yl.year, revenue=yl.gross_revenue,
            opex=yl.total_opex, noi=yl.noi,
            ncf_unlevered=yl.ncf_unlevered, ncf_levered=yl.ncf_levered,
        )
        for yl in pf.years
    ]

    # --- Summary rows ---
    cs0 = pf.contracted_share_schedule[0]
    realized_cf = pf.generation_schedule[0].cf_realized
    summary_rows: list[tuple[str, str]] = [
        ("Technology",            gen.technology.title()),
        ("Market / ISO",          deal.property.market),
        ("Nameplate MW (AC)",     f"{gen.nameplate_mw_ac:,.2f}"),
        ("Target Capacity Factor",_fmt_pct(gen.capacity_factor, 1)),
        ("Realized CF (Yr 1)",    _fmt_pct(realized_cf, 1)),
        ("Net Generation (Yr 1)", f"{pf.years[0].net_generation_mwh:,.0f} MWh"),
        ("Degradation (/yr)",     _fmt_pct(gen.degradation_pct, 2)),
        ("Curtailment",           _fmt_pct(gen.curtailment_pct, 1)),
        ("Availability",          _fmt_pct(gen.availability_pct, 1)),
        ("# Revenue Streams",     str(len(deal.property.revenue_streams))),
        ("Yr-1 Contracted Share", _fmt_pct(cs0.contracted_share, 1)),
        ("All-In Basis / MW",     _fmt_dollar(pf.all_in_basis_per_mw)),
    ]
    if gen.bess_duration_hrs is not None:
        summary_rows.append(("BESS Duration (hrs)", f"{gen.bess_duration_hrs:.1f}"))
        summary_rows.append(("BESS Cycles / Yr", f"{gen.bess_cycles_per_year:.0f}"))
        summary_rows.append(("Round-Trip Efficiency", _fmt_pct(gen.bess_round_trip_eff, 1)))

    # --- Stream roster (Yr 1) ---
    stream_rows: list[tuple[str, str, str, str, str, str, float]] = []
    for stream in deal.property.revenue_streams:
        cpty = f"{stream.counterparty} ({stream.counterparty_rating})"
        # Find the matching per-stream series and its Yr-1 revenue.
        match_key = next(
            (k for k in pf.per_stream_years if k.startswith(stream.label)),
            stream.label,
        )
        yr1 = pf.per_stream_years[match_key][0]
        if isinstance(stream, PPAStream):
            term = f"{stream.start_date.isoformat()} -> {stream.end_date.isoformat()}"
            rate = f"${stream.price_mwh:.2f}/MWh (+{stream.escalation_pct*100:.1f}%/yr)"
            allot = _fmt_pct(stream.allotment_pct, 0)
            kind_label = "PPA"
        elif isinstance(stream, AvailabilityStream):
            term = f"{stream.start_date.isoformat()} -> {stream.end_date.isoformat()}"
            rate = f"${stream.payment_mw_mo:,.0f}/MW-mo on {stream.capacity_mw:.1f} MW"
            allot = "n/a"
            kind_label = "Availability"
        else:  # MerchantStream
            term = f"{stream.market} spot ({len(stream.price_curve_mwh)}-yr curve)"
            rate = f"${stream.price_curve_mwh[0]:.2f}/MWh Yr 1 (+{stream.terminal_growth*100:.1f}% term.)"
            allot = _fmt_pct(stream.allotment_pct, 0)
            kind_label = "Merchant"
        stream_rows.append((
            stream.label, kind_label, cpty, term, rate, allot, yr1.revenue,
        ))

    # --- Contracted vs. merchant mix by year ---
    mix_rows: list[tuple[int, float, float, float, float, float, float]] = []
    for yl, cs in zip(pf.years, pf.contracted_share_schedule):
        total = yl.gross_revenue
        mix_rows.append((
            yl.year, yl.ppa_revenue, yl.availability_revenue,
            yl.merchant_revenue, yl.ptc_revenue, total, cs.contracted_share,
        ))

    # --- Tax credit summary ---
    tc = deal.tax_credits
    tax_credit_rows: list[tuple[str, str]] = []
    if tc.itc_pct > 0 or tc.itc_basis > 0:
        tax_credit_rows.append(("ITC %",       _fmt_pct(tc.itc_pct, 1)))
        tax_credit_rows.append(("ITC Basis",   _fmt_dollar(tc.itc_basis)))
        tax_credit_rows.append(("ITC Cash (Yr 1)",
                                _fmt_dollar(tc.itc_pct * tc.itc_basis)))
    if tc.ptc_per_mwh > 0:
        tax_credit_rows.append(("PTC ($/MWh)", f"${tc.ptc_per_mwh:.2f}"))
        tax_credit_rows.append(("PTC Term (yrs)", str(tc.ptc_term_yrs)))
        tax_credit_rows.append(("PTC Inflation", _fmt_pct(tc.ptc_inflation, 2)))

    block = InfraMemoBlock(
        summary_rows=summary_rows,
        stream_rows=stream_rows,
        mix_rows=mix_rows,
        tax_credit_rows=tax_credit_rows,
    )
    return payload, years, block


def _dispatch(deal_path: str) -> tuple[SummaryPayload, list[MemoYearLine], str, CommercialMemoBlock | None, DCMemoBlock | None, InfraMemoBlock | None]:
    """Detect asset class from YAML, dispatch to correct engine.

    Returns (payload, year_lines, revenue_label, commercial, datacenter, infrastructure).
    """
    import yaml
    with open(deal_path, encoding="utf-8") as f:
        raw = yaml.safe_load(f)
    prop = raw.get("property", {}) or {}
    asset_class = (prop.get("asset_class") or "").lower()
    if not asset_class:
        if "generation" in prop or "revenue_streams" in prop:
            asset_class = "infrastructure"
        elif "keys" in prop or "service_level" in prop or "brand" in prop:
            asset_class = "hospitality"
        elif "unit_mix" in prop:
            asset_class = "multifamily"
        elif "rent_roll" in raw or "leases" in raw:
            asset_class = "office"

    # Datacenter auto-detect fallback if no explicit asset_class
    if not asset_class:
        if "contracts" in prop:
            asset_class = "datacenter_wholesale"
        elif "cabinet_mix" in prop:
            asset_class = "datacenter_colo"

    if asset_class == "multifamily":
        payload, years = _build_mf(deal_path)
        return payload, years, "EGI", None, None, None
    if asset_class in {"office", "industrial", "retail"}:
        payload, years, block = _build_commercial(deal_path)
        return payload, years, "EGI", block, None, None
    if asset_class == "hospitality":
        payload, years = _build_hospitality(deal_path)
        return payload, years, "Total Revenue", None, None, None
    if asset_class == "datacenter_wholesale":
        payload, years, dc_block = _build_dc_wholesale(deal_path)
        return payload, years, "EGI", None, dc_block, None
    if asset_class == "datacenter_colo":
        payload, years, dc_block = _build_dc_colo(deal_path)
        return payload, years, "EGI", None, dc_block, None
    if asset_class == "infrastructure":
        payload, years, infra_block = _build_infrastructure(deal_path)
        return payload, years, "Gross Revenue", None, None, infra_block
    raise ValueError(
        f"unknown asset_class '{asset_class}' in {deal_path}. "
        f"Expected: multifamily / office / industrial / retail / hospitality "
        f"/ datacenter_wholesale / datacenter_colo / infrastructure."
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _print_status(payload: SummaryPayload, out: Path) -> None:
    bar = "-" * 72
    print(bar)
    print(f"  {payload.deal_name}  -  IC Memo")
    print(bar)
    print(f"  Asset class:     {payload.asset_class}")
    print(f"  Purchase Price:  {_fmt_dollar(payload.purchase_price)}")
    print(f"  All-In Basis:    {_fmt_dollar(payload.all_in_basis)}")
    print(f"  LP Net IRR:      {_fmt_pct(payload.lp_irr)}")
    print(f"  GP IRR:          {_fmt_pct(payload.gp_irr)}")
    print(bar)
    print(f"  Memo written to: {out}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m scripts.underwriting.ic_memo",
        description="Generate an institutional IC memo (docx) from a deal YAML.",
    )
    parser.add_argument("deal", help="Path to deal YAML.")
    parser.add_argument(
        "-o", "--output",
        help="Output path (default: outputs/<deal_id>-ic-memo.docx).",
    )
    parser.add_argument(
        "--preparer",
        default="Acquisitions",
        help="Preparer name printed on cover page (default: Acquisitions).",
    )
    args = parser.parse_args(argv)

    payload, years, revenue_label, commercial, datacenter, infrastructure = _dispatch(args.deal)
    out = Path(args.output) if args.output else Path("outputs") / f"{payload.deal_id}-ic-memo.docx"
    write_ic_memo(
        payload, years, out,
        revenue_label=revenue_label, preparer=args.preparer,
        commercial=commercial, datacenter=datacenter,
        infrastructure=infrastructure,
    )
    _print_status(payload, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
